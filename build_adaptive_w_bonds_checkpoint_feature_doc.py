import json
import subprocess
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:/Users/TonyTang/Documents/chan.py")
CHECKPOINT = ROOT / "checkpoints" / "TQQQ_adaptive_reward_fresh_start_252days_at_2020_W_Bonds.joblib"
OUT = ROOT / "TQQQ_adaptive_reward_W_Bonds_checkpoint_daily_and_5min_features_explained.docx"
SYSTEM_PYTHON = Path(r"C:/Users/TonyTang/AppData/Local/Programs/Python/Python313/python.exe")


def daily_category(feature):
    if feature.startswith("k_"):
        return "TQQQ daily kline state"
    if feature.startswith("vix_"):
        return "VIX / volatility context"
    if feature.startswith("ctx_"):
        return "Daily Chan BSP context"
    if feature in {"p", "dp_minK", "dp_maxK"}:
        return "Prior daily probability context"
    if feature.startswith("rg_"):
        return "Daily regime flags"
    if feature.startswith("bd_"):
        return "Base direction flags"
    return "Other"


def daily_explain(feature):
    if feature == "p":
        return "The current/available daily model probability p_day used as recursive model context."
    if feature == "dp_minK":
        return "Current p_day minus recent minimum p_day over the dp_lookback window."
    if feature == "dp_maxK":
        return "Current p_day minus recent maximum p_day over the dp_lookback window."
    if feature == "k_ret1":
        return "TQQQ one-day return."
    if feature.startswith("k_ret_"):
        return f"TQQQ trailing {feature.rsplit('_', 1)[-1]}-day return."
    if feature.startswith("k_vol_"):
        return f"TQQQ trailing {feature.rsplit('_', 1)[-1]}-day realized volatility."
    if feature == "k_atr_14":
        return "TQQQ 14-day ATR; daily volatility scale."
    if feature == "k_range_over_atr":
        return "Current daily range divided by ATR."
    if feature == "k_close_pos":
        return "Close position inside the daily high-low range."
    if feature == "k_gap":
        return "Daily opening gap relative to prior close."
    if feature.startswith("k_above_ma_"):
        return f"Flag: TQQQ close is above its {feature.rsplit('_', 1)[-1]}-day moving average."
    if feature == "k_slope40":
        return "40-day log-price slope for medium-term trend."
    if feature == "vix_level":
        return "Current VIX daily level."
    if feature == "vix_ret1":
        return "One-day VIX return."
    if feature.startswith("vix_ret_"):
        return f"Trailing {feature.rsplit('_', 1)[-1]}-day VIX return."
    if feature.startswith("vix_vol_"):
        return f"Trailing {feature.rsplit('_', 1)[-1]}-day VIX volatility."
    if feature == "vix_atr_14":
        return "VIX 14-day ATR."
    if feature == "vix_range_over_atr":
        return "VIX daily range divided by VIX ATR."
    if feature == "vix_close_pos":
        return "VIX close position inside its daily high-low range."
    if feature == "vix_gap":
        return "VIX opening gap relative to prior close."
    if feature.startswith("vix_above_ma_"):
        return f"Flag: VIX is above its {feature.rsplit('_', 1)[-1]}-day moving average."
    if feature == "vix_slope40":
        return "40-day VIX slope."
    ctx_map = {
        "ctx_has_bsp": "Whether prior daily Chan BSP history exists.",
        "ctx_last_dir_buy": "Flag: latest daily BSP direction is buy.",
        "ctx_last_dir_sell": "Flag: latest daily BSP direction is sell.",
        "ctx_days_since_last_bsp": "Days since the last daily BSP.",
        "ctx_days_since_last_buy": "Days since the last daily buy BSP.",
        "ctx_days_since_last_sell": "Days since the last daily sell BSP.",
        "ctx_density_total": "Count of recent daily BSPs.",
        "ctx_density_buy": "Count of recent daily buy BSPs.",
        "ctx_density_sell": "Count of recent daily sell BSPs.",
        "ctx_density_imb": "Recent BSP direction imbalance.",
        "ctx_price_diff_from_last_buy": "Current close relative to last daily buy BSP close.",
        "ctx_price_diff_from_last_sell": "Current close relative to last daily sell BSP close.",
        "ctx_slope_from_last_buy": "Per-day price slope since last daily buy BSP.",
        "ctx_slope_from_last_sell": "Per-day price slope since last daily sell BSP.",
        "rg_up": "Flag: daily Chan regime is up.",
        "rg_down": "Flag: daily Chan regime is down.",
        "rg_unknown": "Flag: daily Chan regime is unknown.",
        "bd_buy": "Flag: latest base daily BSP direction is buy.",
        "bd_sell": "Flag: latest base daily BSP direction is sell.",
        "bd_none": "Flag: no base daily BSP direction is available.",
    }
    return ctx_map.get(feature, "Daily numeric feature used by the p_day classifier.")


def daily_prediction_meaning(feature):
    cat = daily_category(feature)
    if cat == "TQQQ daily kline state":
        return "Describes trend, volatility, gap, and price location before the daily label is known."
    if cat == "VIX / volatility context":
        return "Adds market risk and volatility regime information."
    if cat == "Daily Chan BSP context":
        return "Shows where today sits relative to recent daily Chan buy/sell points."
    if cat == "Prior daily probability context":
        return "Adds memory of recent p_day behavior."
    if cat == "Daily regime flags":
        return "Marks whether Chan structure is up, down, or unclear."
    if cat == "Base direction flags":
        return "Marks the latest daily BSP direction used for label construction."
    return "Adds context to estimate whether the daily direction will be confirmed."


def daily_leak_note(feature):
    if feature in {"p", "dp_minK", "dp_maxK"}:
        return "Check timing because these are recursive model-state features."
    if feature.startswith("vix_"):
        return "Safe if the VIX daily value is known only after the source-day close."
    if feature.startswith(("ctx_", "rg_", "bd_")):
        return "Safe if built only from daily BSP history known up to the source day."
    return "No direct future target field detected."


def five_category(feature):
    if feature.startswith("prev_daily_vix_"):
        return "5min previous daily VIX context"
    if feature.startswith("prev_daily_us2y_"):
        return "5min previous daily US 2Y context"
    if feature.startswith("prev_daily_us10y_"):
        return "5min previous daily US 10Y context"
    if feature.startswith("prev_daily_yc_"):
        return "5min previous daily yield curve context"
    if feature.startswith("prev_daily_"):
        return "5min previous daily TQQQ context"
    if feature.startswith("feat_bsp") or feature in {
        "feat_bi_amp", "feat_bi_amp_rate", "feat_bi_klu_cnt",
        "feat_break_bi_amp", "feat_break_bi_amp_rate", "feat_break_bi_klu_cnt",
        "feat_bsp_type", "feat_divergence_rate", "feat_level", "feat_zs_cnt", "is_segbsp",
    }:
        return "5min Chan/BSP structure"
    if feature.startswith(("macd", "feat_macd", "rsi", "feat_rsi", "kdj", "feat_kdj", "feat_ppo")):
        return "5min momentum indicators"
    if feature.startswith("dmi"):
        return "5min trend strength indicators"
    if feature.startswith("klu_") or feature in {"body_size", "high_low_spread_pct", "is_bullish_candle", "lower_shadow", "upper_shadow", "price_change_pct"}:
        return "5min current kline state"
    if feature in {"direction_encoded", "bsp_type_encoded", "is_buy"}:
        return "5min signal encoding"
    if feature.startswith("snapshot_"):
        return "5min snapshot bookkeeping"
    return "5min other numeric context"


def five_explain(feature):
    direct = {
        "body_size": "Size of the 5min candle body.",
        "high_low_spread_pct": "Current 5min high-low range as a percent.",
        "upper_shadow": "Upper wick size.",
        "lower_shadow": "Lower wick size.",
        "is_bullish_candle": "Whether the 5min signal candle closed above open.",
        "price_change_pct": "Percent change of the current 5min signal bar.",
        "direction_encoded": "Encoded BSP direction.",
        "is_buy": "Binary buy/sell signal flag.",
        "bsp_type_encoded": "Numeric Chan BSP type encoding.",
        "is_segbsp": "Whether the signal is a segment-level BSP.",
        "feat_bsp_type": "Numeric BSP type from the Chan extractor.",
        "feat_level": "Chan level/timeframe indicator.",
        "feat_zs_cnt": "Number of Zhongshu/central structures.",
        "feat_divergence_rate": "Divergence strength between price and structure/momentum.",
        "feat_ppo": "Percentage Price Oscillator momentum feature.",
        "feat_volume": "Volume feature from the 5min extractor.",
    }
    if feature in direct:
        return direct[feature]
    if feature.startswith("klu_"):
        return f"Raw current 5min K-line {feature.replace('klu_', '')} at the BSP signal."
    if feature.startswith("feat_bsp1"):
        return "Type-1 BSP structure feature."
    if feature.startswith("feat_bsp2s"):
        return "Type-2s BSP structure feature."
    if feature.startswith("feat_bsp2"):
        return "Type-2 BSP structure feature."
    if feature.startswith("feat_bsp3"):
        return "Type-3 BSP structure feature."
    if feature.startswith("feat_bi_"):
        return "Current Bi structure feature."
    if feature.startswith("feat_break_bi_"):
        return "Breakout Bi structure feature."
    if feature.startswith(("macd", "feat_macd")):
        return "MACD momentum feature."
    if feature.startswith(("rsi", "feat_rsi")):
        return "RSI momentum feature."
    if feature.startswith(("kdj", "feat_kdj")):
        return "KDJ oscillator feature."
    if feature.startswith("dmi"):
        return "DMI/ADX trend strength or directional pressure feature."
    if feature.startswith("prev_daily_yc_"):
        return f"Previous completed daily yield-curve feature: {feature.replace('prev_daily_yc_', '').replace('_', ' ')}."
    if feature.startswith("prev_daily_us2y_"):
        return f"Previous completed daily US 2Y feature: {feature.replace('prev_daily_us2y_', '').replace('_', ' ')}."
    if feature.startswith("prev_daily_us10y_"):
        return f"Previous completed daily US 10Y feature: {feature.replace('prev_daily_us10y_', '').replace('_', ' ')}."
    if feature.startswith("prev_daily_vix_"):
        return f"Previous completed daily VIX feature: {feature.replace('prev_daily_vix_', '').replace('_', ' ')}."
    if feature.startswith("prev_daily_"):
        return f"Previous completed daily TQQQ context: {feature.replace('prev_daily_', '').replace('_', ' ')}."
    if feature.startswith("snapshot_"):
        return "BSP snapshot timing/bookkeeping field from signal tracking."
    return "Numeric feature used by the 5min return model."


def five_meaning(feature):
    cat = five_category(feature)
    if cat == "5min signal encoding":
        return "Conditions expected return on signal side/type."
    if cat == "5min current kline state":
        return "Describes immediate price action at the BSP signal."
    if cat == "5min Chan/BSP structure":
        return "Measures whether the BSP has strong Chan structure or is noisy."
    if cat == "5min momentum indicators":
        return "Checks whether momentum confirms or contradicts the BSP."
    if cat == "5min trend strength indicators":
        return "Separates trend setups from choppy conditions."
    if cat.startswith("5min previous daily"):
        return "Adds higher-timeframe context known from the previous completed daily bar."
    if cat == "5min snapshot bookkeeping":
        return "May encode signal timing/order rather than market behavior."
    return "Adds context for estimating forward return after a 5min BSP."


def five_leak_note(feature):
    if feature.startswith("snapshot_"):
        return "Review carefully. Snapshot timing can become a time/order proxy."
    if feature.startswith("prev_daily_"):
        return "No same-day leak if populated only from the previous completed daily bar."
    return "No direct future target field detected from the saved feature name."


def load_checkpoint_data():
    code = rf"""
import json
import joblib
import numpy as np

p = r"{CHECKPOINT}"
b = joblib.load(p)
m = b.get("daily_prob_model")

coef_by_feature = {{}}
classes = []
if m is not None:
    classes = [int(x) for x in list(getattr(m, "classes_", []))]
    for cc in getattr(m, "calibrated_classifiers_", []) or []:
        est = getattr(cc, "estimator", None) or getattr(cc, "base_estimator", None)
        if est is None or not hasattr(est, "named_steps"):
            continue
        vec = est.named_steps.get("vec")
        lr = est.named_steps.get("lr")
        if vec is None or lr is None:
            continue
        names = list(vec.get_feature_names_out())
        coefs = np.asarray(getattr(lr, "coef_", []), dtype=float)
        if coefs.ndim == 2 and coefs.shape[0] >= 1:
            coefs = coefs[0]
        for name, coef in zip(names, coefs):
            coef_by_feature.setdefault(name, []).append(float(coef))

daily_rows = []
for name in sorted(coef_by_feature):
    vals = coef_by_feature[name]
    mean_coef = float(np.mean(vals)) if vals else None
    mean_abs_coef = float(np.mean(np.abs(vals))) if vals else None
    daily_rows.append({{
        "feature": name,
        "mean_coef": mean_coef,
        "mean_abs_coef": mean_abs_coef,
        "fold_count": len(vals),
    }})

buy_pack = b.get("buy_pack") or {{}}
sell_pack = b.get("sell_pack") or {{}}

print(json.dumps({{
    "checkpoint": p,
    "schema": b.get("schema"),
    "code": b.get("code"),
    "snapshot_time": str(b.get("snapshot_time")),
    "N_confirm": b.get("N_confirm"),
    "daily_gate_mode": b.get("daily_gate_mode"),
    "daily_reward_mode": b.get("daily_reward_mode"),
    "daily_model_type": type(m).__name__ if m is not None else None,
    "daily_classes": classes,
    "daily_prob_trained_n": b.get("daily_prob_trained_n"),
    "daily_direct_gate_model": type(b.get("daily_direct_gate_model")).__name__ if b.get("daily_direct_gate_model") is not None else None,
    "daily_direct_gate_trained_n": b.get("daily_direct_gate_trained_n"),
    "daily_rows": daily_rows,
    "five_buy_features": list(buy_pack.get("feature_cols") or []),
    "five_sell_features": list(sell_pack.get("feature_cols") or []),
    "five_buy_model_type": buy_pack.get("model_type"),
    "five_sell_model_type": sell_pack.get("model_type"),
    "lookahead_days_5m": b.get("lookahead_days_5m"),
    "macro_files": b.get("macro_files"),
    "daily_csv_path": b.get("daily_csv_path"),
    "k5m_csv_path": b.get("k5m_csv_path"),
}}, ensure_ascii=True))
"""
    python_exe = str(SYSTEM_PYTHON if SYSTEM_PYTHON.exists() else "python")
    raw = subprocess.check_output([python_exe, "-c", code], cwd=str(ROOT), text=True)
    return json.loads(raw)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=7.6, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_summary_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(["Field", "Value"]):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=8.5, color=(255, 255, 255))
        set_cell_shading(table.rows[0].cells[i], "1F4E79")
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True, size=8.2)
        set_cell_text(cells[1], value, size=8.2)


def add_daily_tables(doc, rows):
    grouped = defaultdict(list)
    for row in rows:
        grouped[daily_category(row["feature"])].append(row)

    order = [
        "TQQQ daily kline state",
        "VIX / volatility context",
        "Daily Chan BSP context",
        "Prior daily probability context",
        "Daily regime flags",
        "Base direction flags",
        "Other",
    ]
    for cat in order:
        items = grouped.get(cat, [])
        if not items:
            continue
        items = sorted(items, key=lambda r: abs(float(r.get("mean_abs_coef") or 0.0)), reverse=True)
        doc.add_heading(cat, level=2)
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Feature", "Mean coef", "Mean abs coef", "Folds", "Logic / calculation", "Meaning in p_day", "Leak note"]
        for i, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], header, bold=True, size=7.8, color=(255, 255, 255))
            set_cell_shading(table.rows[0].cells[i], "1F4E79")
        set_repeat_table_header(table.rows[0])
        for row in items:
            feature = row["feature"]
            cells = table.add_row().cells
            set_cell_text(cells[0], feature, bold=True)
            set_cell_text(cells[1], f"{row.get('mean_coef', 0.0):.6g}")
            set_cell_text(cells[2], f"{row.get('mean_abs_coef', 0.0):.6g}")
            set_cell_text(cells[3], row.get("fold_count", ""))
            set_cell_text(cells[4], daily_explain(feature))
            set_cell_text(cells[5], daily_prediction_meaning(feature))
            set_cell_text(cells[6], daily_leak_note(feature))


def add_five_min_tables(doc, buy_features, sell_features):
    rows = []
    for feature in sorted(set(buy_features) | set(sell_features)):
        used = []
        if feature in buy_features:
            used.append("BUY")
        if feature in sell_features:
            used.append("SELL")
        rows.append({"feature": feature, "used_by": " + ".join(used)})

    grouped = defaultdict(list)
    for row in rows:
        grouped[five_category(row["feature"])].append(row)

    order = sorted(grouped.keys())
    for cat in order:
        doc.add_heading(cat, level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Feature", "Used by", "Logic / calculation", "Meaning in 5min prediction", "Leak note"]
        for i, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], header, bold=True, size=7.8, color=(255, 255, 255))
            set_cell_shading(table.rows[0].cells[i], "1F4E79")
        set_repeat_table_header(table.rows[0])
        for row in grouped[cat]:
            feature = row["feature"]
            cells = table.add_row().cells
            set_cell_text(cells[0], feature, bold=True)
            set_cell_text(cells[1], row["used_by"])
            set_cell_text(cells[2], five_explain(feature))
            set_cell_text(cells[3], five_meaning(feature))
            set_cell_text(cells[4], five_leak_note(feature))


def build_doc(data):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TQQQ Adaptive Reward W_Bonds Checkpoint Features")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.add_run("Checkpoint source: ").bold = True
    p.add_run(data["checkpoint"])

    p = doc.add_paragraph()
    p.add_run("Daily model target: ").bold = True
    p.add_run(
        "binary p_day = P(label=1). label=1 means the current daily Chan base direction is confirmed over "
        "the next N_confirm daily bars: for BUY, future low stays above today's low; for SELL, future high "
        "stays below today's high. The checkpoint then maps p_day into FORCE_BUY / FREE / FORCE_SELL using "
        "threshold logic and prior counterfactual 5min rewards."
    )

    p = doc.add_paragraph()
    p.add_run("5min model target: ").bold = True
    p.add_run(
        "forward return opportunity after a 5min Chan BSP signal, trained separately for BUY and SELL signals."
    )

    add_summary_table(
        doc,
        [
            ("Schema", data.get("schema")),
            ("Code", data.get("code")),
            ("Snapshot time", data.get("snapshot_time")),
            ("N_confirm", data.get("N_confirm")),
            ("Daily model type", data.get("daily_model_type")),
            ("Daily classes", data.get("daily_classes")),
            ("Daily probability training rows", data.get("daily_prob_trained_n")),
            ("Daily gate mode", data.get("daily_gate_mode")),
            ("Daily reward mode", data.get("daily_reward_mode")),
            ("Direct gate model", data.get("daily_direct_gate_model")),
            ("Daily feature count", len(data.get("daily_rows") or [])),
            ("5min BUY feature count", len(data.get("five_buy_features") or [])),
            ("5min SELL feature count", len(data.get("five_sell_features") or [])),
            ("5min BUY model type", data.get("five_buy_model_type")),
            ("5min SELL model type", data.get("five_sell_model_type")),
            ("5min lookahead days", data.get("lookahead_days_5m")),
            ("Macro files", data.get("macro_files")),
        ],
    )

    doc.add_heading("Daily p_day Model Features", level=1)
    add_daily_tables(doc, data.get("daily_rows") or [])

    doc.add_page_break()
    doc.add_heading("5min Return Model Features", level=1)
    add_five_min_tables(doc, data.get("five_buy_features") or [], data.get("five_sell_features") or [])

    doc.add_page_break()
    doc.add_heading("Notebook Verification Snippet", level=1)
    snippet = (
        "import joblib\n\n"
        f"bundle = joblib.load(r\"{CHECKPOINT}\")\n"
        "daily_model = bundle[\"daily_prob_model\"]\n"
        "daily_feature_names = sorted({\n"
        "    name\n"
        "    for cc in daily_model.calibrated_classifiers_\n"
        "    for name in cc.estimator.named_steps[\"vec\"].get_feature_names_out()\n"
        "})\n"
        "buy_5m_features = bundle[\"buy_pack\"][\"feature_cols\"]\n"
        "sell_5m_features = bundle[\"sell_pack\"][\"feature_cols\"]\n"
        "len(daily_feature_names), len(buy_5m_features), len(sell_5m_features)\n"
    )
    para = doc.add_paragraph()
    run = para.add_run(snippet)
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc(load_checkpoint_data()))
