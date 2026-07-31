from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(r"C:/Users/TonyTang/Documents/chan.py")
REFERENCE = Path(r"C:/Users/TonyTang/Desktop/word/5分钟模型.docx")
OUT = ROOT / "日线模型.docx"


FEATURES = [
    "bd_buy",
    "bd_none",
    "bd_sell",
    "ctx_days_since_last_bsp",
    "ctx_days_since_last_buy",
    "ctx_days_since_last_sell",
    "ctx_density_buy",
    "ctx_density_imb",
    "ctx_density_sell",
    "ctx_density_total",
    "ctx_has_bsp",
    "ctx_last_dir_buy",
    "ctx_last_dir_sell",
    "ctx_price_diff_from_last_buy",
    "ctx_price_diff_from_last_sell",
    "ctx_slope_from_last_buy",
    "ctx_slope_from_last_sell",
    "dp_maxK",
    "dp_minK",
    "k_above_ma_100",
    "k_above_ma_20",
    "k_above_ma_50",
    "k_atr_14",
    "k_close_pos",
    "k_gap",
    "k_range_over_atr",
    "k_ret1",
    "k_ret_10",
    "k_ret_20",
    "k_ret_40",
    "k_ret_5",
    "k_slope40",
    "k_vol_10",
    "k_vol_20",
    "k_vol_40",
    "k_vol_5",
    "p",
    "rg_down",
    "rg_unknown",
    "rg_up",
    "vix_above_ma_100",
    "vix_above_ma_20",
    "vix_above_ma_50",
    "vix_atr_14",
    "vix_close_pos",
    "vix_gap",
    "vix_level",
    "vix_range_over_atr",
    "vix_ret1",
    "vix_ret_10",
    "vix_ret_20",
    "vix_ret_40",
    "vix_ret_5",
    "vix_slope40",
    "vix_vol_10",
    "vix_vol_20",
    "vix_vol_40",
    "vix_vol_5",
]


PARAGRAPHS = [
    "日线模型是一个基于日线缠论结构、日线K线状态、VIX波动率环境和历史p_day状态的二分类概率模型。它不直接预测明天的收益，也不直接预测 FORCE_BUY、FORCE_SELL 或 FREE。模型输出的是 p_day，也就是当前日线结构在未来确认窗口内成立的概率。",
    "当前 checkpoint 为 TQQQ_adaptive_reward_fresh_start_252days_at_2020_W_Bonds.joblib。",
    "该 checkpoint 中 daily_prob_model 是 CalibratedClassifierCV，内部使用 DictVectorizer、MaxAbsScaler 和 LogisticRegression。训练样本数 daily_prob_trained_n 为 1825，N_confirm 为 5。",
    "日线模型的标签计算方式为：",
    "如果当前最新日线基础方向 base_dir 为 Buy：",
    "label=1 当且仅当未来 N_confirm 根日线K线中的最低价大于当前日K线最低价。",
    "也就是说，Buy 方向没有在未来确认窗口内跌破当前低点，说明当前买方向结构被确认。",
    "如果当前最新日线基础方向 base_dir 为 Sell：",
    "label=1 当且仅当未来 N_confirm 根日线K线中的最高价小于当前日K线最高价。",
    "也就是说，Sell 方向没有在未来确认窗口内突破当前高点，说明当前卖方向结构被确认。",
    "只有未来 N_confirm 根日线已经完整出现以后，该日期才会获得训练标签。实时预测时只能使用当日收盘及当日之前已经知道的信息。",
    "模型输出 p_day = P(label=1)。之后系统再根据 threshold 逻辑，把 p_day 映射为 FORCE_BUY、FREE 或 FORCE_SELL。",
    "因此在这个版本中，日线模型本身预测的是结构确认概率；gate 是后续策略规则根据 p_day 和历史 counterfactual 5min reward 选择出来的。",
    "一、TQQQ日线K线状态特征",
    "k_ret1：TQQQ当前日线的一日收益率。",
    "k_ret_5：TQQQ过去5日累计收益率。",
    "k_ret_10：TQQQ过去10日累计收益率。",
    "k_ret_20：TQQQ过去20日累计收益率。",
    "k_ret_40：TQQQ过去40日累计收益率。",
    "k_vol_5：TQQQ过去5日收益波动率。",
    "k_vol_10：TQQQ过去10日收益波动率。",
    "k_vol_20：TQQQ过去20日收益波动率。",
    "k_vol_40：TQQQ过去40日收益波动率。",
    "k_atr_14：TQQQ 14日ATR，用来衡量日线波动尺度。",
    "k_range_over_atr：当前日线最高价到最低价的范围除以ATR，表示当日波动相对正常波动是否异常。",
    "k_close_pos：当前收盘价在当日最高价和最低价区间中的位置。",
    "k_gap：当前开盘价相对前一日收盘价的跳空幅度。",
    "k_above_ma_20：当前收盘价是否高于20日均线。",
    "k_above_ma_50：当前收盘价是否高于50日均线。",
    "k_above_ma_100：当前收盘价是否高于100日均线。",
    "k_slope40：过去40日对数价格斜率，用来描述中期趋势方向和强度。",
    "二、VIX波动率环境特征",
    "vix_level：当前VIX收盘水平，表示市场风险和恐慌程度。",
    "vix_ret1：VIX一日变化率。",
    "vix_ret_5：VIX过去5日变化率。",
    "vix_ret_10：VIX过去10日变化率。",
    "vix_ret_20：VIX过去20日变化率。",
    "vix_ret_40：VIX过去40日变化率。",
    "vix_vol_5：VIX过去5日波动率。",
    "vix_vol_10：VIX过去10日波动率。",
    "vix_vol_20：VIX过去20日波动率。",
    "vix_vol_40：VIX过去40日波动率。",
    "vix_atr_14：VIX 14日ATR，衡量波动率本身的波动尺度。",
    "vix_range_over_atr：VIX当日区间除以VIX ATR。",
    "vix_close_pos：VIX收盘价在当日高低区间中的位置。",
    "vix_gap：VIX开盘相对前一日收盘的跳空幅度。",
    "vix_above_ma_20：VIX是否高于20日均线。",
    "vix_above_ma_50：VIX是否高于50日均线。",
    "vix_above_ma_100：VIX是否高于100日均线。",
    "vix_slope40：VIX过去40日趋势斜率。",
    "当前 checkpoint 的 macro_files 包含 VIX、US2Y 和 US10Y，但实际 fitted daily p_day model 的 feature names 中只出现了 vix_* 宏观变量，没有出现 us2y_* 或 us10y_*。",
    "三、日线缠论BSP上下文特征",
    "ctx_has_bsp：当前日期之前是否已经出现过日线BSP。",
    "ctx_last_dir_buy：最近一个日线BSP方向是否为Buy。",
    "ctx_last_dir_sell：最近一个日线BSP方向是否为Sell。",
    "ctx_days_since_last_bsp：距离最近任意方向日线BSP经过的天数。",
    "ctx_days_since_last_buy：距离最近Buy方向日线BSP经过的天数。",
    "ctx_days_since_last_sell：距离最近Sell方向日线BSP经过的天数。",
    "ctx_density_total：最近一段窗口内日线BSP总数量。",
    "ctx_density_buy：最近一段窗口内Buy方向日线BSP数量。",
    "ctx_density_sell：最近一段窗口内Sell方向日线BSP数量。",
    "ctx_density_imb：最近BSP方向不平衡度，反映Buy/Sell结构密度偏向。",
    "ctx_price_diff_from_last_buy：当前收盘价相对最近Buy BSP价格的变化。",
    "ctx_price_diff_from_last_sell：当前收盘价相对最近Sell BSP价格的变化。",
    "ctx_slope_from_last_buy：从最近Buy BSP到当前收盘价的平均每日价格斜率。",
    "ctx_slope_from_last_sell：从最近Sell BSP到当前收盘价的平均每日价格斜率。",
    "这些特征描述当前日线位置与历史缠论买卖点之间的距离、方向和结构密度。",
    "四、历史p_day状态特征",
    "p：当前可用的日线模型概率值。训练时作为递归状态特征进入后续样本。",
    "dp_minK：当前p_day减去最近dp_lookback窗口内p_day最小值。",
    "dp_maxK：当前p_day减去最近dp_lookback窗口内p_day最大值。",
    "这组特征描述模型概率自身的短期变化，例如p_day是否正在从低位抬升，或是否已经从高位回落。",
    "五、日线趋势区域和基础方向特征",
    "rg_up：当前日线缠论结构被判断为上行区域。",
    "rg_down：当前日线缠论结构被判断为下行区域。",
    "rg_unknown：当前日线结构无法明确判断为上行或下行。",
    "bd_buy：最新日线基础BSP方向为Buy。",
    "bd_sell：最新日线基础BSP方向为Sell。",
    "bd_none：当前没有可用的日线基础方向。",
    "base direction 既参与标签定义，也作为模型特征帮助模型理解当前结构背景。",
    "六、与gate的关系",
    "日线模型先输出 p_day。",
    "系统随后在 daily_gate_mode=threshold 下，根据历史 daily_reward_log 中 FORCE_BUY、FREE、FORCE_SELL 的 counterfactual 5min reward，动态选择 buy_level 和 sell_level。",
    "如果 p_day 落入不同阈值区间，系统才会决定当天使用 FORCE_BUY、FREE 或 FORCE_SELL。",
    "因此这个模型不是直接三分类gate模型，而是gate系统中的日线概率输入。",
    "七、未来信息和泄漏风险说明",
    "label 使用未来 N_confirm 日线来计算，但这个未来信息只用于训练标签。实时预测时不会使用未来K线。",
    "k_* 特征只能使用当前日线及历史日线数据。",
    "vix_* 特征只有在对应日期收盘后已经可知时，才适合用于下一交易日决策。",
    "ctx_*、rg_* 和 bd_* 应只基于当前日期之前或当前日期已经确认的缠论结构。",
    "p、dp_minK 和 dp_maxK 是递归模型状态特征，需要特别检查时间顺序，确保它们来自当时已经产生的p_day，而不是未来重新计算后的结果。",
    "未使用字段包括：",
    "未来最高价、未来最低价、未来确认结果、best_action_ex_post、reward_force_buy、reward_free、reward_force_sell 等未来结果字段都不应作为模型训练特征。",
    "八、当前实际使用的全部58个特征",
]

PARAGRAPHS.extend(FEATURES)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_paragraph(paragraph, text):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def main():
    doc = Document(str(REFERENCE))
    body = doc._body._element
    paragraphs = doc.paragraphs

    while len(paragraphs) < len(PARAGRAPHS):
        doc.add_paragraph()
        paragraphs = doc.paragraphs

    for i, text in enumerate(PARAGRAPHS):
        set_paragraph(paragraphs[i], text)

    for paragraph in paragraphs[len(PARAGRAPHS):]:
        body.remove(paragraph._element)

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
