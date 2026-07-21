import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:/Users/TonyTang/Documents/chan.py")
FEATURE_IMPORTANCE = ROOT / "feature_importance.csv"
FEATURES_USED = ROOT / "features_used_in_training.csv"
OUT = ROOT / "5min_model_training_variables_explanation.docx"


def read_features():
    features = []
    importance = {}
    if FEATURE_IMPORTANCE.exists():
        with FEATURE_IMPORTANCE.open("r", encoding="utf-8-sig", newline="") as f:
            for row in csv.DictReader(f):
                feat = (row.get("feature") or "").strip()
                if not feat:
                    continue
                features.append(feat)
                try:
                    importance[feat] = float(row.get("importance", ""))
                except Exception:
                    importance[feat] = None
    elif FEATURES_USED.exists():
        with FEATURES_USED.open("r", encoding="utf-8-sig", newline="") as f:
            for row in csv.DictReader(f):
                feat = (row.get("feature_name") or "").strip()
                if feat:
                    features.append(feat)

    seen = set()
    ordered = []
    for feat in features:
        if feat not in seen:
            seen.add(feat)
            ordered.append(feat)
    return ordered, importance


EXCLUDED_BY_CURRENT_SELECTOR = {
    "timestamp",
    "code",
    "direction",
    "bsp_type",
    "klu_idx",
    "best_return_pct",
    "has_best_exit",
    "best_exit_type",
    "best_exit_klu_idx",
    "best_exit_price",
    "regime_bucket",
}

CATEGORY_ORDER = [
    "疑似 label / target / 未来信息",
    "K线价格与成交量",
    "短期收益与价格变化",
    "均线与趋势位置",
    "动量 / 超买超卖指标",
    "波动率与趋势强度",
    "成交量结构",
    "K线形态",
    "价格形态 / 支撑阻力",
    "BSP / Chan 结构特征",
    "编码与方向变量",
    "其他",
]


def horizon(v):
    match = re.search(r"_(\d+)$", v)
    return match.group(1) if match else ""


def category(v):
    if v.startswith("target_return_") or v.startswith("label_") or v in {
        "feat_next_bi_return",
        "profit_target_distance",
    }:
        return "疑似 label / target / 未来信息"
    if v.startswith("klu_") or v in {
        "price_change_pct",
        "high_low_spread_pct",
        "upper_shadow",
        "lower_shadow",
        "body_size",
        "is_bullish_candle",
    }:
        return "K线价格与成交量"
    if v.startswith("return_"):
        return "短期收益与价格变化"
    if v.startswith(("sma_", "ema_", "price_above_sma_", "price_above_ema_")):
        return "均线与趋势位置"
    if v.startswith(
        (
            "macd",
            "feat_macd",
            "rsi",
            "feat_rsi",
            "kdj",
            "feat_kdj",
            "stoch",
            "roc",
            "tsi",
            "feat_ppo",
            "cci",
            "williams",
            "uo",
            "mfi",
        )
    ):
        return "动量 / 超买超卖指标"
    if v.startswith(("atr", "dmi", "psar", "price_above_psar")):
        return "波动率与趋势强度"
    if v.startswith("volume_") or v == "feat_volume":
        return "成交量结构"
    if v.startswith("candle_"):
        return "K线形态"
    if v.startswith("price_"):
        return "价格形态 / 支撑阻力"
    if v.startswith("feat_bsp") or v.startswith("feat_zs") or v == "feat_divergence_rate":
        return "BSP / Chan 结构特征"
    if v in {
        "is_buy",
        "direction",
        "direction_encoded",
        "bsp_type",
        "bsp_types",
        "bsp_type_encoded",
        "has_profit_target",
    }:
        return "编码与方向变量"
    return "其他"


def logic(v):
    if v in EXCLUDED_BY_CURRENT_SELECTOR:
        return "当前代码会作为 identifier / metadata 排除，不应进入最终 X 矩阵。"
    if v.startswith("target_return_") or v.startswith("label_") or v in {
        "feat_next_bi_return",
        "profit_target_distance",
    }:
        return "用于历史标注或结果分析的字段；如果在实时训练特征中出现，需要检查 future leak。"
    if v == "klu_open":
        return "当前 5min K线开盘价，提供信号发生时的价格位置。"
    if v == "klu_high":
        return "当前 5min K线最高价，反映短周期上冲强度。"
    if v == "klu_low":
        return "当前 5min K线最低价，反映短周期下探压力。"
    if v == "klu_close":
        return "当前 5min K线收盘价，是信号确认时的基准价格。"
    if v == "klu_volume":
        return "当前 5min K线成交量，衡量信号是否有成交确认。"
    if v.startswith("return_"):
        return f"过去 {horizon(v)} 根左右K线的短期收益，用来判断已有 momentum 或短线反转。"
    if v == "price_change_pct":
        return "当前K线涨跌幅，衡量本根K线方向和力度。"
    if v == "high_low_spread_pct":
        return "当前K线高低价振幅，衡量日内短周期波动。"
    if v == "upper_shadow":
        return "上影线比例，反映冲高回落压力。"
    if v == "lower_shadow":
        return "下影线比例，反映下探后承接力量。"
    if v == "body_size":
        return "K线实体大小，衡量多空单边推动力度。"
    if v == "is_bullish_candle":
        return "是否阳线，表示该5分钟bar收盘是否强于开盘。"
    if v.startswith("sma_"):
        return f"{horizon(v)}周期简单均线，表示短中期价格中枢。"
    if v.startswith("ema_"):
        return f"{horizon(v)}周期指数均线，对近期价格变化更敏感，用来识别趋势位置。"
    if v.startswith("price_above_sma_"):
        return f"价格是否在 SMA{horizon(v)} 上方，用来判断趋势偏多或偏空。"
    if v.startswith("price_above_ema_"):
        return f"价格是否在 EMA{horizon(v)} 上方，用来判断近期趋势方向。"
    if v in {
        "macd_value",
        "macd_dif",
        "macd_dea",
        "macd_signal",
        "feat_macd_value",
        "feat_macd_dea",
        "feat_macd_diff",
    }:
        return "MACD相关变量，用来衡量趋势动能、快慢线差异和动能变化。"
    if v == "feat_ppo":
        return "PPO动量指标，类似百分比形式的MACD，用来比较不同价格水平下的趋势力度。"
    if v in {"rsi", "feat_rsi", "rsi_oversold", "rsi_overbought"}:
        return "RSI相关变量，用来识别短周期超买、超卖和反转风险。"
    if v.startswith(("kdj", "feat_kdj")):
        return "KDJ相关变量，用来衡量短周期摆动、超买超卖和拐点。"
    if v.startswith("stoch"):
        return "Stochastic指标，用来判断价格在近期区间中的相对位置。"
    if v.startswith("roc_"):
        return f"ROC{horizon(v.replace('_positive', ''))}动量/方向变量，衡量价格变化速度。"
    if v.startswith("tsi"):
        return "TSI趋势强度指标，用来衡量动量是否持续。"
    if v.startswith("cci"):
        return "CCI指标，用来识别价格偏离常态区间的程度。"
    if v.startswith("williams"):
        return "Williams %R 指标，用来识别短线超买或超卖。"
    if v.startswith("uo"):
        return "Ultimate Oscillator，用多周期动量判断短线反转风险。"
    if v.startswith("mfi"):
        return "Money Flow Index，结合价格和成交量判断资金流入/流出。"
    if v == "atr":
        return "ATR波动率，衡量近期价格波动幅度。"
    if v == "atr_ratio":
        return "ATR相对价格比例，用来比较不同价格水平下的波动强弱。"
    if v.startswith("dmi"):
        return "DMI/ADX趋势强度变量，用来区分趋势行情和震荡行情。"
    if v == "psar" or v == "price_above_psar":
        return "Parabolic SAR相关变量，用来判断趋势跟踪止损/反转位置。"
    if v.startswith("volume_") or v == "feat_volume":
        return "成交量结构变量，用来判断信号是否伴随放量、缩量、吸筹或派发。"
    if v.startswith("candle_"):
        name = v.replace("candle_", "").replace("_", " ")
        return f"K线形态变量：{name}，用于捕捉短线反转或延续形态。"
    if v.startswith("price_"):
        name = v.replace("price_", "").replace("_", " ")
        return f"价格结构变量：{name}，用于识别突破、整理、支撑/阻力或趋势形态。"
    if v == "feat_divergence_rate":
        return "背驰/背离程度，衡量价格新高新低与动能之间是否不一致。"
    if v.startswith("feat_bsp1"):
        return "1类买卖点相关结构特征，描述第一类BSP附近笔的幅度、长度和力度。"
    if v.startswith("feat_bsp2s"):
        return "2s类买卖点结构特征，描述次级二类点的回撤、突破和笔结构。"
    if v.startswith("feat_bsp2"):
        return "2类买卖点结构特征，描述回撤、突破、笔幅度和笔数量。"
    if v.startswith("feat_bsp3"):
        return "3类买卖点结构特征，描述中枢高度、离开中枢后的笔幅度和持续性。"
    if v == "feat_bsp_bi_amp":
        return "当前BSP对应笔的幅度，用来衡量该信号结构的强弱。"
    if v == "feat_zs_cnt":
        return "中枢数量，表示当前结构复杂度和震荡程度。"
    if v in {"is_buy", "direction_encoded"}:
        return "信号方向编码，用来区分buy信号和sell信号的预测逻辑。"
    if v in {"bsp_type_encoded", "bsp_types"}:
        return "BSP类型编码，用来让模型区分不同类型买卖点的历史胜率和收益分布。"
    if v == "has_profit_target":
        return "是否存在目标收益/止盈信息；若由未来结果生成，需要避免作为实时输入。"
    return "数值型训练变量，用来补充模型对5min信号状态的判断。"


def meaning(v):
    cat = category(v)
    if cat == "疑似 label / target / 未来信息":
        return "这类字段能解释历史表现，但不应在实时预测时依赖；需要确认是否只用于标签、评估或离线分析。"
    if cat == "BSP / Chan 结构特征":
        return "帮助模型判断当前买卖点是否是高质量结构，而不是单纯价格噪音。"
    if cat == "动量 / 超买超卖指标":
        return "帮助模型判断短线动能是否支持该 buy/sell signal，以及是否存在反转风险。"
    if cat == "均线与趋势位置":
        return "帮助模型判断信号发生在趋势顺势位置还是逆势位置。"
    if cat == "波动率与趋势强度":
        return "帮助模型识别趋势是否足够强，以及当前波动是否会放大信号风险。"
    if cat == "成交量结构":
        return "帮助模型判断价格变化是否有成交量支持。"
    if cat == "K线形态":
        return "帮助模型识别短线反转、延续和多空力量变化。"
    if cat == "价格形态 / 支撑阻力":
        return "帮助模型判断信号是否发生在突破、整理、支撑或阻力附近。"
    if cat == "K线价格与成交量":
        return "提供信号发生时最基础的5分钟市场状态。"
    if cat == "短期收益与价格变化":
        return "反映信号前的短线价格路径，用来判断追涨、反转或延续。"
    if cat == "编码与方向变量":
        return "让模型区分不同方向和不同BSP类型的收益分布。"
    return "作为补充数值信息，提升模型对5min信号质量的区分能力。"


def set_font(run, size=8.5, bold=False):
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


def main():
    features, importance = read_features()
    features_sorted = sorted(
        features,
        key=lambda x: (
            CATEGORY_ORDER.index(category(x)) if category(x) in CATEGORY_ORDER else 999,
            features.index(x),
        ),
    )

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    styles["Normal"].font.size = Pt(9)
    for sty in ["Heading 1", "Heading 2"]:
        styles[sty].font.name = "Arial"
        styles[sty]._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("5min Model Training Variables List"), size=18, bold=True)

    p = doc.add_paragraph()
    set_font(p.add_run("说明："), bold=True)
    set_font(
        p.add_run(
            "该文档根据项目中的 5min model 训练流程整理。当前代码中，5min model 使用 buy/sell point rows "
            "训练两个 XGBoost return regressors，目标变量是 best_return_pct；特征筛选逻辑来自 "
            "pipelineCurrent.py 的 prepare_ml_dataset 和 get_feature_columns。"
        )
    )

    p = doc.add_paragraph()
    set_font(p.add_run("注意："), bold=True)
    set_font(
        p.add_run(
            "表中标记为“疑似 label / target / 未来信息”的变量，如果出现在训练特征里，需要重点检查是否会造成 "
            "future leak；它们更适合作为标签、评估或离线分析字段，而不是实时预测输入。"
        )
    )

    summary = doc.add_table(rows=1, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.style = "Table Grid"
    summary.rows[0].cells[0].text = "来源"
    summary.rows[0].cells[1].text = "内容"
    for row in [
        ("Feature importance", "feature_importance.csv"),
        ("Training feature list", "features_used_in_training.csv"),
        ("Feature selector", "pipelineCurrent.py: prepare_ml_dataset / get_feature_columns"),
        ("变量数量", str(len(features_sorted))),
    ]:
        cells = summary.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]

    doc.add_paragraph("")

    current_cat = None
    table = None
    for feat in features_sorted:
        cat = category(feat)
        if cat != current_cat:
            doc.add_heading(cat, level=2)
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            headers = ["Variable", "Importance", "是否当前选择器排除", "逻辑 / 计算含义", "在5min预测中的意义"]
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_font(run, bold=True)
            current_cat = cat

        cells = table.add_row().cells
        cells[0].text = feat
        imp = importance.get(feat)
        cells[1].text = "" if imp is None else f"{imp:.6g}"
        cells[2].text = (
            "是"
            if feat in EXCLUDED_BY_CURRENT_SELECTOR
            else ("需检查" if cat == "疑似 label / target / 未来信息" else "否")
        )
        cells[3].text = logic(feat)
        cells[4].text = meaning(feat)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_font(run)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
